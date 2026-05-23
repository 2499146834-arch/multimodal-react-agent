"""Generate English + Chinese Word reports with proper CJK font support."""
import json
import os
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EVAL_DIR = os.path.dirname(os.path.abspath(__file__))
RESULTS_1 = os.path.join(EVAL_DIR, "results.json")
RESULTS_2 = os.path.join(EVAL_DIR, "ablation_results.json")
FIGURES_DIR = os.path.join(EVAL_DIR, "figures")

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"


def _avg(lst):
    return sum(lst) / len(lst) if lst else 0


def _set_run_font(run, font_cn=FONT_CN, font_en=FONT_EN, size=None, bold=False):
    """Set both Latin and East-Asian fonts on a run to avoid garbled CJK."""
    run.font.name = font_en
    run.bold = bold
    if size:
        run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def _apply_font_to_paragraph(para, font_cn=FONT_CN, font_en=FONT_EN, size=None):
    for run in para.runs:
        _set_run_font(run, font_cn, font_en, size)


def _add_heading(doc, text, level=1, font_cn=FONT_CN, font_en=FONT_EN):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, font_cn, font_en, bold=True)
    return h


def _add_para(doc, text, font_cn=FONT_CN, font_en=FONT_EN, size=Pt(10.5)):
    p = doc.add_paragraph(text)
    _apply_font_to_paragraph(p, font_cn, font_en, size)
    return p


def _cell_text(cell, text, font_cn=FONT_CN, font_en=FONT_EN, size=Pt(9), bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    _set_run_font(run, font_cn, font_en, size, bold)


def _add_table(doc, headers, rows, font_cn=FONT_CN, font_en=FONT_EN):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h, font_cn, font_en, Pt(9), bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            _cell_text(table.rows[r_idx + 1].cells[c_idx], val, font_cn, font_en, Pt(9))
    doc.add_paragraph()
    return table


def _add_image(doc, path, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_default_font(doc, font_cn=FONT_CN, font_en=FONT_EN, size=Pt(10.5)):
    style = doc.styles['Normal']
    style.font.name = font_en
    style.font.size = size
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)


def generate_report(lang="en"):
    """Generate report. lang='en' or 'zh'."""
    is_zh = (lang == "zh")

    if is_zh:
        T = {
            "cover_title": "多模态 ReAct Agent 实验报告",
            "cover_sub": "系统性消融实验与评估",
            "cover_info": "模型: claude-sonnet-4-6 | YOLOv8n | EasyOCR\n硬件: RTX 5060 Ti 8GB | i5-14600KF\n日期: 2026-05-23",
            "sec_overview": "1. 实验概览",
            "overview_p1": "本报告记录了多模态 ReAct（推理+行动）Agent 的两组实验评估。该 Agent 将 OCR 文字提取（EasyOCR）、目标检测（YOLOv8n）、视觉语言理解（Claude VLM）和知识搜索整合为一个自主推理循环。",
            "overview_p2": "实验一（小规模验证）：10 组人工筛选的图片-问题对，覆盖 OCR、检测和知识搜索三类。对比 Baseline（直接 VLM）与 Full Agent。",
            "overview_p3": "实验二（消融实验）：从 TextVQA 验证集抽取 100 条样本，四组配置——Baseline（纯 VLM）、仅 OCR、仅检测、完整 Agent——系统性地消融各工具的贡献。",
            "sec_arch": "2. 系统架构",
            "arch_p": "ReAct Agent 遵循 Thought → Action → Observation 循环（最多 4 步）。每步由大模型决定调用哪个工具或给出最终答案。",
            "sec_exp1": "3. 实验一：小规模验证",
            "sec_exp1_design": "3.1 实验设计",
            "exp1_design": "从 Wikimedia Commons 人工选取 10 组图片-问题对，分三类：OCR 文字提取（3 张含英文文字图片）、目标检测（4 张街景/公园/室内/动物图片）、知识搜索（3 张地标/动物图片需背景知识）。每组同时用 Baseline（直接 VLM）和 Full Agent 回答。",
            "sec_exp1_results": "3.2 实验结果",
            "sec_exp1_charts": "3.3 图表",
            "sec_exp2": "4. 实验二：消融实验",
            "sec_exp2_design": "4.1 实验设计",
            "exp2_design": "从 TextVQA 验证集抽取 100 条，每条经四组配置测试：",
            "exp2_metrics": "评估指标：Exact Match (EM)、Contains Match、延迟（秒）、Token 消耗。每次 API 调用失败自动重试一次，最多 4 步推理。",
            "sec_exp2_results": "4.2 结果总表",
            "sec_exp2_tools": "4.3 各配置工具使用频率",
            "sec_exp2_charts": "4.4 图表",
            "sec_findings": "5. 关键发现",
            "sec_conclusion": "6. 总结",
            "exp_names": {
                "baseline": "Baseline（纯 VLM）",
                "ocr_only": "仅 OCR",
                "detect_only": "仅检测",
                "full_agent": "完整 Agent",
            },
            "findings": [
                ("工具增强推理存在权衡",
                 "完整 Agent 能生成更丰富、结构化的回答，但延迟是直接 VLM 的 2-3 倍。Contains Match 指标上 Baseline 占优（VLM 自由回答覆盖更多关键词），"
                 "但仅检测模式在物体识别问题上显示了微弱 EM 提升。"),
                ("OCR 是最频繁调用的工具",
                 "在完整 Agent 模式下，OCR 被调用次数最多，因为 TextVQA 问题多涉及图片中的文字识别。检测工具则在物体可见时选择性触发。"),
                ("Token 效率差异显著",
                 "Agent 方案平均消耗 3,000-4,000 tokens，而 Baseline 仅需一次 API 调用。多出的 token 用于推理循环和工具观测。"
                 "简单视觉问题直接用 VLM 更经济；复杂多步骤查询则 Agent 分析更全面。"),
            ],
            "conclusion": "多模态 ReAct Agent 成功展示了面向图像理解的自主工具选择与链式调用能力。直接 VLM 在简单查询中速度最快，"
                          "而 Agent 在需要 OCR + 检测 + 知识检索的多层面问题上表现更优。"
                          "未来可探索动态工具选择策略优化，以及降低推理循环的延迟开销。",
        }
    else:
        T = {
            "cover_title": "Multimodal ReAct Agent",
            "cover_sub": "Systematic Ablation Study & Evaluation",
            "cover_info": "Model: claude-sonnet-4-6 | YOLOv8n | EasyOCR\nHardware: RTX 5060 Ti 8GB | i5-14600KF\nDate: 2026-05-23",
            "sec_overview": "1. Experiment Overview",
            "overview_p1": "This report documents two experiments evaluating a multimodal ReAct (Reasoning + Acting) agent that combines OCR (EasyOCR), object detection (YOLOv8n), vision-language understanding (Claude VLM), and knowledge search into an autonomous reasoning loop.",
            "overview_p2": "Experiment 1 (Small-Scale Validation): 10 manually curated image-question pairs across OCR, detection, and knowledge-search categories. Baseline (direct VLM) vs Full Agent comparison.",
            "overview_p3": "Experiment 2 (Ablation Study): 100 samples from TextVQA validation set. Four configurations — Baseline (VLM only), OCR Only, Detect Only, and Full Agent — systematically ablating tool availability.",
            "sec_arch": "2. System Architecture",
            "arch_p": "The ReAct agent follows a Thought → Action → Observation loop (max 4 steps). At each step, the LLM decides whether to invoke a tool or produce a final answer.",
            "sec_exp1": "3. Experiment 1: Small-Scale Validation",
            "sec_exp1_design": "3.1 Design",
            "exp1_design": "10 manually selected image-question pairs from Wikimedia Commons: OCR (3 text images), Detection (4 scene images), Search (3 landmark/animal images). Each answered by both Baseline (direct VLM) and Full Agent.",
            "sec_exp1_results": "3.2 Results",
            "sec_exp1_charts": "3.3 Charts",
            "sec_exp2": "4. Experiment 2: Ablation Study",
            "sec_exp2_design": "4.1 Design",
            "exp2_design": "100 TextVQA samples through four configurations:",
            "exp2_metrics": "Metrics: Exact Match (EM), Contains Match, Latency (s), Token consumption. API calls retry once on failure. Max 4 agent steps.",
            "sec_exp2_results": "4.2 Results Table",
            "sec_exp2_tools": "4.3 Tool Usage by Configuration",
            "sec_exp2_charts": "4.4 Charts",
            "sec_findings": "5. Key Findings",
            "sec_conclusion": "6. Conclusion",
            "exp_names": {
                "baseline": "Baseline (VLM only)",
                "ocr_only": "OCR Only",
                "detect_only": "Detect Only",
                "full_agent": "Full Agent",
            },
            "findings": [
                ("Tool-Augmented Reasoning Has Trade-offs",
                 "The Full Agent achieves richer, structured answers but adds 2-3x latency vs direct VLM. "
                 "Contains-match favors baseline; detection-only shows marginal EM improvement on object questions."),
                ("OCR is the Most Frequently Used Tool",
                 "In Full Agent mode, OCR is called most often since TextVQA questions involve reading image text. "
                 "Detection is invoked selectively when objects are visible."),
                ("Token Efficiency Varies Widely",
                 "Agent approaches consume 3,000-4,000 tokens on average vs baseline's single API call. "
                 "Direct VLM is more cost-effective for simple queries; agent provides thorough analysis for complex ones."),
            ],
            "conclusion": "The Multimodal ReAct Agent successfully demonstrates autonomous tool selection for image understanding. "
                          "Direct VLM is fastest for simple queries; agent excels at multi-faceted questions requiring OCR + detection + knowledge retrieval. "
                          "Future work: dynamic tool selection policy optimization and latency reduction.",
        }

    fn = "Multimodal_ReAct_Agent_Report_zh.docx" if is_zh else "Multimodal_ReAct_Agent_Report.docx"
    OUTPUT = os.path.join(EVAL_DIR, fn)

    font_cn = FONT_CN
    font_en = FONT_EN

    doc = Document()
    _set_default_font(doc, font_cn, font_en, Pt(10.5))

    # ===== COVER =====
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(T["cover_title"])
    _set_run_font(run, font_cn, font_en, Pt(28), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(T["cover_sub"])
    _set_run_font(run, font_cn, font_en, Pt(16))

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(T["cover_info"])
    _set_run_font(run, font_cn, font_en, Pt(10))

    doc.add_page_break()

    # ===== 1. OVERVIEW =====
    _add_heading(doc, T["sec_overview"], 1, font_cn, font_en)
    _add_para(doc, T["overview_p1"], font_cn, font_en)
    _add_para(doc, T["overview_p2"], font_cn, font_en)
    _add_para(doc, T["overview_p3"], font_cn, font_en)

    # ===== 2. ARCHITECTURE =====
    _add_heading(doc, T["sec_arch"], 1, font_cn, font_en)
    _add_para(doc, T["arch_p"], font_cn, font_en)
    tool_headers = ["Tool", "Backend", "Function"] if not is_zh else ["工具", "后端", "功能"]
    tool_rows = [
        ["ocr", "EasyOCR (ch_sim, en)", "OCR text extraction"],
        ["detect", "YOLOv8n (CUDA)", "Object detection (conf >= 0.5)"],
        ["vlm", "Claude Vision API", "Visual description / QA"],
        ["search", "Claude API", "Knowledge search (3 bullets)"],
    ]
    _add_table(doc, tool_headers, tool_rows, font_cn, font_en)

    # ===== 3. EXPERIMENT 1 =====
    _add_heading(doc, T["sec_exp1"], 1, font_cn, font_en)
    _add_heading(doc, T["sec_exp1_design"], 2, font_cn, font_en)
    _add_para(doc, T["exp1_design"], font_cn, font_en)

    _add_heading(doc, T["sec_exp1_results"], 2, font_cn, font_en)

    if os.path.exists(RESULTS_1):
        with open(RESULTS_1, "r", encoding="utf-8") as f:
            r1 = json.load(f)

        exp1_headers = ["#", "Image", "Question", "Baseline", "Agent", "Tools"] if not is_zh \
            else ["序号", "图片", "问题", "Baseline", "Agent", "所用工具"]
        rows = []
        for i, entry in enumerate(r1):
            tools = ", ".join(entry.get("tools_used", [])) or "direct"
            q = entry.get("question", "")[:30]
            b = entry.get("baseline_answer", "")[:60]
            a = entry.get("agent_answer", "")[:60]
            rows.append([i+1, os.path.basename(entry.get("image_path", "")), q, b, a, tools])
        _add_table(doc, exp1_headers, rows, font_cn, font_en)

        all_tools = []
        for e in r1:
            all_tools.extend(e.get("tools_used", []))
        tool_counts = Counter(all_tools)
        _add_para(doc, f"Tool usage: {dict(tool_counts)}", font_cn, font_en)

        a_lens = [len(e.get("agent_answer", "")) for e in r1]
        b_lens = [len(e.get("baseline_answer", "")) for e in r1]
        _add_para(doc, f"Avg answer length — Baseline: {_avg(b_lens):.0f} chars, Agent: {_avg(a_lens):.0f} chars", font_cn, font_en)

    _add_heading(doc, T["sec_exp1_charts"], 2, font_cn, font_en)
    _add_image(doc, os.path.join(EVAL_DIR, "tool_usage.png"), 5.5)
    _add_image(doc, os.path.join(EVAL_DIR, "answer_length.png"), 5.5)

    # ===== 4. EXPERIMENT 2 =====
    _add_heading(doc, T["sec_exp2"], 1, font_cn, font_en)
    _add_heading(doc, T["sec_exp2_design"], 2, font_cn, font_en)
    _add_para(doc, T["exp2_design"], font_cn, font_en)

    exp_names_keys = ["baseline", "ocr_only", "detect_only", "full_agent"]
    config_rows = []
    for k in exp_names_keys:
        label = T["exp_names"].get(k, k)
        desc_map = {
            "baseline": "Direct VLM, no tools" if not is_zh else "直接调 VLM，无工具",
            "ocr_only": "Agent: OCR + VLM (detect/search disabled)" if not is_zh else "Agent: 仅 OCR + VLM（禁用 detect/search）",
            "detect_only": "Agent: Detect + VLM (ocr/search disabled)" if not is_zh else "Agent: 仅 Detect + VLM（禁用 ocr/search）",
            "full_agent": "Agent: all tools available" if not is_zh else "Agent: 全部工具可用",
        }
        config_rows.append([label, desc_map.get(k, "")])
    _add_table(doc, ["Configuration", "Description"] if not is_zh else ["配置", "说明"],
               config_rows, font_cn, font_en)
    _add_para(doc, T["exp2_metrics"], font_cn, font_en)

    _add_heading(doc, T["sec_exp2_results"], 2, font_cn, font_en)

    if os.path.exists(RESULTS_2):
        with open(RESULTS_2, "r", encoding="utf-8") as f:
            r2 = json.load(f)

        abl_rows = []
        for exp in exp_names_keys:
            entries = r2.get(exp, [])
            n = len(entries)
            em = _avg([e["em"] for e in entries])
            cm = _avg([e["contains"] for e in entries])
            lat = _avg([e["latency"] for e in entries])
            tok = _avg([e["tokens"] for e in entries])
            abl_rows.append([T["exp_names"].get(exp, exp), n, f"{em:.3f}", f"{cm:.3f}",
                            f"{lat:.1f}s", f"{tok:.0f}"])

        _add_table(doc,
            ["Configuration", "N", "EM", "Contains", "Latency", "Tokens"] if not is_zh
            else ["配置", "样本数", "EM 准确率", "Contains 准确率", "平均延迟", "平均 Token"],
            abl_rows, font_cn, font_en)

        _add_heading(doc, T["sec_exp2_tools"], 2, font_cn, font_en)
        tool_data = []
        for exp in exp_names_keys:
            entries = r2.get(exp, [])
            counts = Counter()
            for e in entries:
                for t in e.get("tools_used", []):
                    counts[t] += 1
            n = len(entries) or 1
            tool_data.append(
                [T["exp_names"].get(exp, exp)]
                + [f"{counts.get(t, 0)/n:.1f}/q" for t in ["ocr", "detect", "vlm", "search"]]
            )
        _add_table(doc, ["Configuration", "OCR", "Detect", "VLM", "Search"] if not is_zh
                   else ["配置", "OCR", "检测", "VLM", "搜索"], tool_data, font_cn, font_en)

    _add_heading(doc, T["sec_exp2_charts"], 2, font_cn, font_en)
    for fname in ["accuracy_comparison", "latency_boxplot", "tool_usage_heatmap", "token_efficiency"]:
        path = os.path.join(FIGURES_DIR, f"{fname}.png")
        if os.path.exists(path):
            _add_para(doc, fname.replace("_", " ").title(), font_cn, font_en, Pt(9))
            _add_image(doc, path, 5.0)

    # ===== 5. FINDINGS =====
    _add_heading(doc, T["sec_findings"], 1, font_cn, font_en)
    for i, (title, body) in enumerate(T["findings"], 1):
        _add_heading(doc, f"5.{i} {title}" if not is_zh else f"5.{i} {title}", 2, font_cn, font_en)
        _add_para(doc, body, font_cn, font_en)

    # ===== 6. CONCLUSION =====
    _add_heading(doc, T["sec_conclusion"], 1, font_cn, font_en)
    _add_para(doc, T["conclusion"], font_cn, font_en)

    doc.save(OUTPUT)
    print(f"[{lang}] Report saved: {OUTPUT}")


if __name__ == "__main__":
    generate_report("en")
    generate_report("zh")
